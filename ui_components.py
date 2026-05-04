import sys
import json
import os
import docx
from PyQt5.QtWidgets import (QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QSplitter, QScrollArea, QLabel, 
                             QTextEdit, QFileDialog, QLineEdit, QCheckBox, 
                             QMessageBox, QInputDialog, QProgressBar)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QPixmap, QImage, QTextCursor, QTextDocument
from PyQt5.QtPrintSupport import QPrinter

from document_handler import DocumentHandler
from ocr_engine import OCREngine
from search_util import SearchUtil, Trie

class OCRWorker(QThread):
    progress = pyqtSignal(int)
    result = pyqtSignal(int, str)
    finished = pyqtSignal()

    def __init__(self, document_handler, ocr_engine):
        super().__init__()
        self.doc_handler = document_handler
        self.ocr_engine = ocr_engine
        self.is_running = True

    def run(self):
        total_pages = self.doc_handler.num_pages
        for i in range(total_pages):
            if not self.is_running:
                break
            img = self.doc_handler.get_page(i)
            text = self.ocr_engine.process_image(img) if img else ""
            self.result.emit(i, text)
            self.progress.emit(i + 1)
        self.finished.emit()

    def stop(self):
        self.is_running = False

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Transliterated-Search Tool")
        self.resize(1200, 800)

        self.doc_handler = DocumentHandler()
        self.ocr_engine = OCREngine()
        
        self.page_texts = []
        self.current_page = 0
        self.total_pages = 0
        self.current_search_results = []
        self.current_search_index = -1

        self.worker = None

        self.init_ui()
        self.load_config()

    def init_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QVBoxLayout(main_widget)

        # ---- Top Toolbar ----
        toolbar_layout = QHBoxLayout()
        
        self.btn_open = QPushButton("Open PDF/Image")
        self.btn_open.setMinimumWidth(180)
        self.btn_open.clicked.connect(self.open_file)

        self.btn_load_model = QPushButton("Load Custom Model")
        self.btn_load_model.setMinimumWidth(180)
        self.btn_load_model.clicked.connect(self.load_model)
        
        self.btn_ocr = QPushButton("Run OCR")
        self.btn_ocr.setMinimumWidth(100)
        self.btn_ocr.clicked.connect(self.run_ocr)
        self.btn_ocr.setEnabled(False)

        self.btn_save_text = QPushButton("Save Text")
        self.btn_save_text.setMinimumWidth(100)
        self.btn_save_text.clicked.connect(self.save_text)
        self.btn_save_text.setEnabled(False)

        self.btn_load_text = QPushButton("Load Text")
        self.btn_load_text.setMinimumWidth(100)
        self.btn_load_text.clicked.connect(self.load_text)
        self.btn_load_text.setEnabled(False)

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search OCR Text (Use ITRANS e.g. bhaarat)")
        
        self.chk_transliterate = QCheckBox("Transliterate")
        self.chk_transliterate.setChecked(True)

        self.chk_fuzzy = QCheckBox("Fuzzy Search")
        self.chk_fuzzy.setChecked(False)

        self.btn_search = QPushButton("Search")
        self.btn_search.clicked.connect(self.perform_search)

        self.btn_prev = QPushButton("<")
        self.btn_prev.clicked.connect(self.search_prev)

        self.btn_next = QPushButton(">")
        self.btn_next.clicked.connect(self.search_next)

        toolbar_layout.addWidget(self.btn_open)
        toolbar_layout.addWidget(self.btn_load_model)
        toolbar_layout.addWidget(self.btn_ocr)
        toolbar_layout.addWidget(self.btn_save_text)
        toolbar_layout.addWidget(self.btn_load_text)
        toolbar_layout.addSpacing(20)
        toolbar_layout.addWidget(self.search_input)
        toolbar_layout.addWidget(self.chk_transliterate)
        toolbar_layout.addWidget(self.chk_fuzzy)
        toolbar_layout.addWidget(self.btn_search)
        toolbar_layout.addWidget(self.btn_prev)
        toolbar_layout.addWidget(self.btn_next)
        
        # ---- Navigation Toolbar ----
        nav_layout = QHBoxLayout()
        self.btn_page_prev = QPushButton("< Prev Page")
        self.btn_page_prev.clicked.connect(lambda: self.navigate_to_page(self.current_page - 1))
        self.lbl_page_info = QLabel("Page 0 / 0")
        self.lbl_page_info.setAlignment(Qt.AlignCenter)
        self.btn_page_next = QPushButton("Next Page >")
        self.btn_page_next.clicked.connect(lambda: self.navigate_to_page(self.current_page + 1))
        
        nav_layout.addStretch()
        nav_layout.addWidget(self.btn_page_prev)
        nav_layout.addWidget(self.lbl_page_info)
        nav_layout.addWidget(self.btn_page_next)
        nav_layout.addStretch()

        # ---- Main Content Splitter ----
        self.splitter = QSplitter(Qt.Horizontal)
        
        # Left Panel (Document)
        self.left_scroll = QScrollArea()
        self.left_scroll.setWidgetResizable(True)
        self.image_label = QLabel("PDF Viewer Area")
        self.image_label.setAlignment(Qt.AlignCenter)
        self.left_scroll.setWidget(self.image_label)

        # Right Panel (OCR Text)
        self.text_edit = QTextEdit()
        self.text_edit.setPlaceholderText("OCR Text Area")

        self.splitter.addWidget(self.left_scroll)
        self.splitter.addWidget(self.text_edit)
        self.splitter.setSizes([600, 600])

        main_layout.addLayout(toolbar_layout)
        main_layout.addLayout(nav_layout)
        main_layout.addWidget(self.splitter)
        
        # Setup Status bar and Progress bar
        self.statusBar()
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.statusBar().addPermanentWidget(self.progress_bar, 1)

    def load_model(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Tesseract Model (.traineddata)", "", "Tesseract Model (*.traineddata)")
        if file_path:
            tessdata_dir = os.path.dirname(file_path)
            model_name = os.path.basename(file_path).split('.')[0]
            self.ocr_engine.set_config(lang=model_name, tessdata_dir=tessdata_dir)
            
            config_data = {"lang": model_name, "tessdata_dir": tessdata_dir}
            try:
                with open("config.json", "w") as f:
                    json.dump(config_data, f)
            except Exception as e:
                print(f"Failed to save config: {e}")

            QMessageBox.information(self, "Model Loaded", f"Successfully loaded custom OCR model: {model_name}\n\nConfiguration has been saved for future sessions.")
            self.btn_ocr.setText(f"Run OCR")

    def load_config(self):
        config_path = "config.json"
        if os.path.exists(config_path):
            try:
                with open(config_path, "r") as f:
                    config_data = json.load(f)
                    lang = config_data.get("lang")
                    tessdata_dir = config_data.get("tessdata_dir")
                    if lang and tessdata_dir and os.path.exists(tessdata_dir):
                        self.ocr_engine.set_config(lang=lang, tessdata_dir=tessdata_dir)
                        self.btn_ocr.setText(f"Run OCR")
            except Exception as e:
                print(f"Failed to load config: {e}")

    def open_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open Document", "", "Supported Files (*.pdf *.png *.jpg *.jpeg *.bmp *.tiff)")
        if file_path:
            try:
                self.total_pages = self.doc_handler.load_document(file_path)
                self.page_texts = [""] * self.total_pages
                self.current_search_results.clear()
                self.current_search_index = -1
                
                self.btn_ocr.setEnabled(True)
                self.btn_save_text.setEnabled(True)
                self.btn_load_text.setEnabled(True)
                
                self.navigate_to_page(0)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load document: {e}")

    def navigate_to_page(self, index):
        if self.total_pages == 0:
            return
            
        # Save current text from UI to memory
        if 0 <= self.current_page < self.total_pages:
            self.page_texts[self.current_page] = self.text_edit.toPlainText()
            
        if index < 0:
            index = 0
        elif index >= self.total_pages:
            index = self.total_pages - 1
            
        self.current_page = index
        self.lbl_page_info.setText(f"Page {self.current_page + 1} / {self.total_pages}")
        
        self.btn_page_prev.setEnabled(self.current_page > 0)
        self.btn_page_next.setEnabled(self.current_page < self.total_pages - 1)
        
        # Load image
        img = self.doc_handler.get_page(index)
        if img:
            data = img.tobytes("raw", "RGB")
            bytes_per_line = 3 * img.size[0]
            qim = QImage(data, img.size[0], img.size[1], bytes_per_line, QImage.Format_RGB888).copy()
            pixmap = QPixmap.fromImage(qim)
            scaled_pixmap = pixmap.scaledToWidth(550, Qt.SmoothTransformation)
            self.image_label.setPixmap(scaled_pixmap)
        else:
            self.image_label.clear()
            self.image_label.setText("Failed to load image")
            
        # Load text
        self.text_edit.setPlainText(self.page_texts[self.current_page])

    def run_ocr(self):
        if self.total_pages == 0:
            return

        langs = self.ocr_engine.get_available_languages()
        if not langs:
            langs = ['eng']
            
        current_lang = self.ocr_engine.lang
        try:
            current_idx = langs.index(current_lang)
        except ValueError:
            current_idx = 0
            if current_lang not in langs:
                langs.insert(0, current_lang)

        lang, ok = QInputDialog.getItem(self, "Select OCR Language", "Select language for OCR:", langs, current_idx, False)
        if not ok or not lang:
            return 

        self.ocr_engine.set_config(lang=lang, tessdata_dir=self.ocr_engine.tessdata_dir)

        self.btn_ocr.setEnabled(False)
        self.btn_open.setEnabled(False)
        
        self.progress_bar.setRange(0, self.total_pages)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(True)

        self.worker = OCRWorker(self.doc_handler, self.ocr_engine)
        self.worker.progress.connect(self.progress_bar.setValue)
        self.worker.result.connect(self.update_ocr_text)
        self.worker.finished.connect(self.ocr_finished)
        self.worker.start()

    def update_ocr_text(self, index, text):
        if 0 <= index < self.total_pages:
            self.page_texts[index] = text
            if index == self.current_page:
                self.text_edit.setPlainText(text)

    def cancel_ocr(self):
        if self.worker:
            self.worker.stop()

    def ocr_finished(self):
        self.btn_ocr.setEnabled(True)
        self.btn_open.setEnabled(True)
        self.progress_bar.setVisible(False)

    def save_text(self):
        if self.total_pages == 0:
            QMessageBox.warning(self, "No Text", "There is no document to save.")
            return
            
        # Ensure current page text is committed
        self.page_texts[self.current_page] = self.text_edit.toPlainText()
        
        file_path, selected_filter = QFileDialog.getSaveFileName(
            self, "Save OCR Text", "", "Text Files (*.txt);;Word Documents (*.docx);;PDF Files (*.pdf);;All Files (*)"
        )
        if file_path:
            try:
                if file_path.endswith('.docx') or "Word Document" in selected_filter:
                    doc = docx.Document()
                    for i, text in enumerate(self.page_texts):
                        doc.add_heading(f'Page {i+1}', level=1)
                        doc.add_paragraph(text)
                    if not file_path.endswith('.docx'):
                        file_path += '.docx'
                    doc.save(file_path)
                    QMessageBox.information(self, "Success", "Text saved successfully as DOCX.")
                    
                elif file_path.endswith('.pdf') or "PDF File" in selected_filter:
                    printer = QPrinter(QPrinter.HighResolution)
                    printer.setOutputFormat(QPrinter.PdfFormat)
                    if not file_path.endswith('.pdf'):
                        file_path += '.pdf'
                    printer.setOutputFileName(file_path)
                    
                    doc = QTextDocument()
                    html_content = ""
                    for i, text in enumerate(self.page_texts):
                        html_content += f"<h1>Page {i+1}</h1><pre style='font-family: monospace;'>{text}</pre><br>"
                    doc.setHtml(html_content)
                    doc.print_(printer)
                    QMessageBox.information(self, "Success", "Text saved successfully as PDF.")
                    
                else:
                    if not file_path.endswith('.txt') and "Text File" in selected_filter:
                        file_path += '.txt'
                    with open(file_path, "w", encoding="utf-8") as f:
                        for i, text in enumerate(self.page_texts):
                            f.write(f"--- Page {i+1} ---\n")
                            f.write(text)
                            f.write("\n\n")
                    QMessageBox.information(self, "Success", "Text saved successfully as TXT.")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save text: {e}")

    def load_text(self):
        if self.total_pages == 0:
            QMessageBox.warning(self, "No Document", "Please load a PDF or Image first.")
            return

        file_path, _ = QFileDialog.getOpenFileName(
            self, "Load OCR Text", "", "Text/Word Files (*.txt *.docx);;All Files (*)"
        )
        if not file_path:
            return

        try:
            import re
            texts = []
            if file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    parts = re.split(r'^--- Page \d+ ---$', content, flags=re.MULTILINE)
                    for part in parts[1:]:
                        texts.append(part.strip())
            elif file_path.endswith('.docx'):
                doc = docx.Document(file_path)
                current_page_text = []
                for p in doc.paragraphs:
                    if p.style.name.startswith('Heading') and p.text.startswith('Page '):
                        if current_page_text:
                            texts.append('\n'.join(current_page_text).strip())
                            current_page_text = []
                    else:
                        current_page_text.append(p.text)
                if current_page_text:
                    texts.append('\n'.join(current_page_text).strip())
            
            if not texts:
                QMessageBox.warning(self, "Warning", "Could not find any parsed pages in the file. Ensure it was saved by this tool.")
                return

            for i in range(min(self.total_pages, len(texts))):
                self.page_texts[i] = texts[i]
            
            self.text_edit.setPlainText(self.page_texts[self.current_page])
            QMessageBox.information(self, "Success", f"Loaded text for {len(texts)} pages.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load text: {e}")

    def perform_search(self):
        if self.total_pages == 0:
            return
            
        # commit current page text
        self.page_texts[self.current_page] = self.text_edit.toPlainText()
            
        self.current_search_results.clear()
        self.current_search_index = -1
        
        # remove highlighted text in active edit
        cursor = self.text_edit.textCursor()
        cursor.select(QTextCursor.Document)
        fmt = cursor.charFormat()
        fmt.setBackground(Qt.transparent)
        cursor.setCharFormat(fmt)
        cursor.clearSelection()
        self.text_edit.setTextCursor(cursor)
            
        raw_query = self.search_input.text()
        if not raw_query.strip():
            return

        query = raw_query
        if self.chk_transliterate.isChecked():
            query = SearchUtil.transliterate_query(raw_query)
            self.search_input.setText(f"{raw_query} => {query}")
        else:
            if "=>" in raw_query:
                query = raw_query.split("=>")[1].strip()

        words_to_search = {query}

        if self.chk_fuzzy.isChecked():
            trie = Trie()
            for text in self.page_texts:
                for token in text.split():
                    cleaned = token.strip('.,()[]{}"\'`~?!;:-\n\t')
                    if cleaned:
                        trie.insert(cleaned)
            
            fuzzy_matches = set()
            for q_word in query.split():
                matches = trie.search(q_word)
                for m, _ in matches:
                    fuzzy_matches.add(m)
            
            if fuzzy_matches:
                words_to_search = fuzzy_matches

        # Find occurrences across all page texts
        for i, text in enumerate(self.page_texts):
            doc = QTextDocument()
            doc.setPlainText(text)
            for w in words_to_search:
                cursor = QTextCursor(doc)
                while not cursor.isNull() and not cursor.atEnd():
                    if self.chk_fuzzy.isChecked():
                        cursor = doc.find(w, cursor)
                    else:
                        cursor = doc.find(w, cursor, QTextDocument.FindCaseSensitively)
                        
                    if not cursor.isNull():
                        self.current_search_results.append((i, cursor.selectionStart(), cursor.selectionEnd()))

        self.current_search_results.sort(key=lambda x: (x[0], x[1]))

        if self.current_search_results:
            self.current_search_index = 0
            self.highlight_current_search()
        else:
            QMessageBox.information(self, "Search", f"No results found for '{query}'.")

    def highlight_current_search(self):
        if not self.current_search_results or self.current_search_index < 0:
            return
            
        page_idx, start_pos, end_pos = self.current_search_results[self.current_search_index]
        
        # navigate if needed
        if self.current_page != page_idx:
            self.navigate_to_page(page_idx)
            
        # highlight
        cursor = self.text_edit.textCursor()
        cursor.select(QTextCursor.Document)
        fmt = cursor.charFormat()
        fmt.setBackground(Qt.transparent)
        cursor.setCharFormat(fmt)
        
        cursor.setPosition(start_pos)
        cursor.setPosition(end_pos, QTextCursor.KeepAnchor)
        
        fmt = cursor.charFormat()
        fmt.setBackground(Qt.yellow)
        cursor.setCharFormat(fmt)
        
        self.text_edit.setTextCursor(cursor)
        self.text_edit.setFocus()

    def search_next(self):
        if not self.current_search_results:
            return
        self.current_search_index = (self.current_search_index + 1) % len(self.current_search_results)
        self.highlight_current_search()

    def search_prev(self):
        if not self.current_search_results:
            return
        self.current_search_index = (self.current_search_index - 1) % len(self.current_search_results)
        self.highlight_current_search()
